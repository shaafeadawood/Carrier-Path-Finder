import io
import docx
import PyPDF2
from fastapi import UploadFile

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from uploaded CV files (PDF or DOCX).
    
    Args:
        file (UploadFile): The uploaded file
        
    Returns:
        str: Extracted text content
    """
    print(f"Processing file: {file.filename}, content type: {file.content_type}")
    content = await file.read()
    text = ""
    
    # Process based on file type
    if file.filename.lower().endswith('.pdf'):
        # PDF processing
        try:
            print("Processing PDF file...")
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            print(f"PDF has {len(pdf_reader.pages)} pages")
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                print(f"Page {i+1} extracted: {len(page_text)} characters")
                text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            text = ""
    
    elif file.filename.lower().endswith('.docx'):
        # DOCX processing
        try:
            print("Processing DOCX file...")
            doc = docx.Document(io.BytesIO(content))
            print(f"DOCX has {len(doc.paragraphs)} paragraphs")
            text = "\n".join([para.text for para in doc.paragraphs])
            print(f"Extracted {len(text)} characters from DOCX")
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            text = ""
    else:
        print(f"Unsupported file type: {file.filename}")
        
    # Reset file cursor for potential future use
    await file.seek(0)
    
    print(f"Final extracted text length: {len(text)}")
    if len(text) > 100:
        print(f"Text preview: {text[:100]}...")
    else:
        print(f"Full text: {text}")
    
    return text
